import pandas as pd
import docx.oxml.ns as ns
from docx import Document
import re
from docx.shared import RGBColor
import docx.oxml as oxml
import openai

API_KEY_FILE = 'apikey.txt'

# 1. read csv file
def read_csv_file(file_path):
    try:
        # Read the CSV file into a DataFrame
        df = pd.read_csv(file_path)
        # Keep only the specified fields (columns)
        fields = ["Title", "Funder", "Deadline", "Amount", "Eligibility", "Abstract", "More Information"]
        filtered_df = df[fields]
        # Return the DataFrame
        return filtered_df
    except FileNotFoundError:
        print("File not found!")
        return None
    except KeyError:
        print("Some specified fields do not exist in the CSV file!")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

# 2. convert the csv file to word file
def df_to_word(data_frame, file_path):
    # Create a new Word document
    doc = Document()

    # Add a table to the document with the same number of columns as the DataFrame
    table = doc.add_table(rows=1, cols=len(data_frame.columns))

    # Get the cells of the header row in the table
    hdr_cells = table.rows[0].cells

    # Fill in the header row with the column names of the DataFrame
    for i, column_name in enumerate(data_frame.columns):
        hdr_cells[i].text = column_name

    # Loop over each row in the DataFrame
    for index, row in data_frame.iterrows():
        # Add a new row to the table
        row_cells = table.add_row().cells
        # Fill in the cells of the row with the values from the DataFrame
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Save the Word document to the specified file path
    doc.save(file_path)

# 3. format the word file
def format_word_file(input_file_path, output_file_path):
    get_api_key()
    # Load the Word document
    doc = Document(input_file_path)

    # Create a new Word document for the formatted content
    formatted_doc = Document()

    # Loop over each table in the document
    for table in doc.tables:
        # Loop over each row in the table, skipping the first row
        for row_index, row in enumerate(table.rows[1:]):
            cells = row.cells
            # Create a string to store the formatted text for the row
            formatted_text = ""
            more_information_link = ""
            title_text = ""
            for cell_index, cell in enumerate(cells):
                # Use the cell text as the value
                value = cell.text
                # Remove HTML tags
                value = re.sub(r"<[^>]*>", "", value)
                header = table.cell(0, cell_index).text
                # Append the label and value to the formatted text string
                if header == "Title":
                    title_text = value
                elif header == "Deadline":
                    formatted_text += f"Due Date: {value}\n"
                elif header == "Amount":
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",  # Specify the chat model
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                            {"role": "user", "content": f"Summarize the following text by extracting award amount in"
                                                        f"USD. Is amount upper exists, just use that number is enough."
                                                        f"Do not include any notes or explanations:\n\n{value}"}
                        ],
                        max_tokens=150,  # Set the maximum length for the summary
                        temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                    )
                    # The response format is different for chat completions
                    summary = response['choices'][0]['message']['content'].strip()
                    print(f'extracted amount {summary}')
                    formatted_text += f"Award Amount: {summary}\n"
                elif header == "Eligibility":
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",  # Specify the chat model
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                            {"role": "user", "content": f"Summarize the following text by extracting which level "
                                                        f"of faculty is eligible. If the level is not mentioned, "
                                                        f"simply return Any level faculty:\n\n{value}"}
                        ],
                        max_tokens=150,  # Set the maximum length for the summary
                        temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                    )
                    # The response format is different for chat completions
                    summary = response['choices'][0]['message']['content'].strip()
                    formatted_text += f"Eligibility: {summary}\n"
                elif header == "Abstract":
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",  # Specify the chat model
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant who's good at summarization."},
                            {"role": "user", "content": f"Summarize the following text in a concise way:\n\n{value}"}
                        ],
                        max_tokens=150,  # Set the maximum length for the summary
                        temperature=0.7  # Adjusts randomness in the response. Lower is more deterministic.
                    )
                    # The response format is different for chat completions
                    summary = response['choices'][0]['message']['content'].strip()
                    formatted_text += f"Program Goal: {summary}\n"
                elif header == "More Information":
                    more_information_link = value
            # Add the formatted text to the new Word document
            p = formatted_doc.add_paragraph()
            if more_information_link:
                hyperlink_run = add_hyperlink(p, more_information_link, title_text)
                hyperlink_run.font.color.rgb = RGBColor(0, 0, 255)
                hyperlink_run.bold = True
            else:
                title_run = p.add_run(title_text)
                title_run.bold = True
            p.add_run("\n" + formatted_text.strip())

    # Save the formatted Word document
    formatted_doc.save(output_file_path)

# Helper functions:
# Add hyperlink
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = oxml.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.ns.qn('r:id'), r_id)

    new_run = oxml.OxmlElement('w:r')
    rPr = oxml.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)

    return r

def get_api_key():
    with open(API_KEY_FILE, 'r') as file:
        openai.api_key = file.read().strip()


if __name__ == "__main__":
    # 1. read csv file
    file_path = "sample_data/opps_export.csv"
    data_frame = read_csv_file(file_path)
    if data_frame is not None:
        print(data_frame.head())  # Print first 5 rows of the DataFrame

        # 2. convert csv file to word file
        word_file_path = "output_word/output.docx"
        df_to_word(data_frame, word_file_path)

        # 3. format word file
        formatted_word_file_path = "output_word/formattedOutput.docx"
        format_word_file(word_file_path, formatted_word_file_path)